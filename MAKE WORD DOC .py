import tkinter as tk
from tkinter import filedialog
from docx import Document

def create_new_document():
    global doc
    doc = Document()
    update_status("New document created.")

def add_paragraph():
    text = entry.get()
    if doc and text:
        doc.add_paragraph(text)
        update_status("Paragraph added.")
    else:
        update_status("No document or no text!")

def save_document():
    if doc:
        file_path = filedialog.asksaveasfilename(defaultextension=".docx")
        if file_path:
            doc.save(file_path)
            update_status(f"Document saved as {file_path}.")
    else:
        update_status("No document to save.")

def update_status(message):
    status_label.config(text=message)

# Set up the main window
root = tk.Tk()
root.title("Word Document Creator")

# Create widgets
create_doc_button = tk.Button(root, text="Create New Document", command=create_new_document)
add_para_button = tk.Button(root, text="Add Paragraph", command=add_paragraph)
save_doc_button = tk.Button(root, text="Save Document", command=save_document)
entry = tk.Entry(root, width=50)
status_label = tk.Label(root, text="", fg="blue")

# Arrange widgets
create_doc_button.pack(pady=5)
add_para_button.pack(pady=5)
entry.pack(pady=5)
save_doc_button.pack(pady=5)
status_label.pack(pady=5)

# Start the GUI event loop
root.mainloop()
