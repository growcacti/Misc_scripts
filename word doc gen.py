import tkinter as tk
from tkinter import filedialog
from docx import Document

# Function to add a paragraph to the document
def add_paragraph():
    paragraph_text = text_entry.get("1.0", tk.END).strip()  # Get text from the text entry
    if paragraph_text:  # If there's text, add a paragraph to the document
        document.add_paragraph(paragraph_text)
        text_entry.delete("1.0", tk.END)  # Clear the text entry field

# Function to save the document
def save_document():
    path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx"), ("All files", "*.*")])
    if path:  # If a file path is given, save the document
        document.save(path)

# Create a new Word document
document = Document()

# Create the main window
root = tk.Tk()
root.title("Word Document Generator")

# Add a Text widget for entering text
text_entry = tk.Text(root, height=10, width=50)
text_entry.pack()

# Add a button to add a paragraph to the document
add_button = tk.Button(root, text="Add Paragraph", command=add_paragraph)
add_button.pack()

# Add a button to save the document
save_button = tk.Button(root, text="Save Document", command=save_document)
save_button.pack()

# Start the GUI event loop
root.mainloop()
